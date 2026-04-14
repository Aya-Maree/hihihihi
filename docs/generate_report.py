"""
Generates Reliability_and_Evaluation_Report.docx from the report content.
Run from the docs/ directory: python generate_report.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.8)
section.right_margin  = Cm(2.8)

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1A, 0x3A, 0x5C)   # headings
MID_BLUE   = RGBColor(0x2E, 0x6D, 0xA8)   # sub-headings / accents
LIGHT_GREY = RGBColor(0xF2, 0xF4, 0xF7)   # table header bg (set via shading)
TEXT_GREY  = RGBColor(0x44, 0x44, 0x44)   # body text
CODE_BG    = RGBColor(0xF5, 0xF5, 0xF5)

# ── Helper: shade a table cell ─────────────────────────────────────────────────
def shade_cell(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

# ── Helper: set paragraph style safely ────────────────────────────────────────
def safe_style(para, style_name: str):
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        pass

# ── Helper: add a run with formatting ─────────────────────────────────────────
def add_run(para, text: str, bold=False, italic=False,
            color: RGBColor = None, size_pt: int = None, font_name: str = None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size_pt:
        run.font.size = Pt(size_pt)
    if font_name:
        run.font.name = font_name
    return run

# ── Helper: paragraph with spacing ────────────────────────────────────────────
def add_para(text="", bold=False, italic=False, color: RGBColor = None,
             size_pt=11, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6, font_name="Calibri"):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if text:
        add_run(p, text, bold=bold, italic=italic,
                color=color, size_pt=size_pt, font_name=font_name)
    return p

# ── Helper: heading ────────────────────────────────────────────────────────────
def add_heading(text: str, level: int = 1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if level == 1:
        pf.space_before = Pt(18)
        pf.space_after  = Pt(6)
        add_run(p, text, bold=True, color=DARK_BLUE, size_pt=16, font_name="Calibri")
    elif level == 2:
        pf.space_before = Pt(12)
        pf.space_after  = Pt(4)
        add_run(p, text, bold=True, color=MID_BLUE, size_pt=13, font_name="Calibri")
    elif level == 3:
        pf.space_before = Pt(8)
        pf.space_after  = Pt(3)
        add_run(p, text, bold=True, color=DARK_BLUE, size_pt=11, font_name="Calibri")
    return p

# ── Helper: horizontal rule ────────────────────────────────────────────────────
def add_hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E6DA8")
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── Helper: bullet point ───────────────────────────────────────────────────────
def add_bullet(text: str, indent_level=0, bold_prefix: str = ""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.25 + indent_level * 0.25)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, color=TEXT_GREY, size_pt=10.5, font_name="Calibri")
        add_run(p, text, color=TEXT_GREY, size_pt=10.5, font_name="Calibri")
    else:
        add_run(p, text, color=TEXT_GREY, size_pt=10.5, font_name="Calibri")
    return p

# ── Helper: code block ─────────────────────────────────────────────────────────
def add_code(lines: list):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        add_run(p, line, size_pt=9, font_name="Courier New", color=RGBColor(0x22, 0x22, 0x22))

# ── Helper: simple table ───────────────────────────────────────────────────────
def add_table(headers: list, rows: list, col_widths: list = None):
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        shade_cell(cell, "1A3A5C")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size  = Pt(10)
        run.font.name  = "Calibri"

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = t.rows[r_idx + 1]
        bg  = "F2F4F7" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            run.font.color.rgb = TEXT_GREY

    # Column widths
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing after table


# ══════════════════════════════════════════════════════════════════════════════
# COVER / TITLE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()  # top padding

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(4)
add_run(title_p, "EventOps AI", bold=True, color=DARK_BLUE, size_pt=28, font_name="Calibri")

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(4)
add_run(sub_p, "Reliability and Evaluation Report", bold=False, color=MID_BLUE, size_pt=16, font_name="Calibri")

course_p = doc.add_paragraph()
course_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
course_p.paragraph_format.space_after = Pt(2)
add_run(course_p, "SE4471B Course Project  |  Assignment 3", color=TEXT_GREY, size_pt=11, font_name="Calibri")

authors_p = doc.add_paragraph()
authors_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors_p.paragraph_format.space_after = Pt(2)
add_run(authors_p, "Sara Daher  ·  Aya Maree", bold=True, color=TEXT_GREY, size_pt=11, font_name="Calibri")

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.paragraph_format.space_after = Pt(0)
add_run(date_p, "April 2026", color=TEXT_GREY, size_pt=11, font_name="Calibri")

add_hr()
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — SYSTEM OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading("1. System Overview", 1)
add_para(
    "EventOps AI is a conversational AI event planning system that transforms unstructured natural-language "
    "event requirements into a complete, conflict-checked planning package. The final system achieves "
    "Tier 2 (Tool-Enabled System), incorporating:",
    color=TEXT_GREY, size_pt=11
)

bullets_overview = [
    ("RAG pipeline: ", "sentence-transformers (all-MiniLM-L6-v2) + ChromaDB over a 23-document knowledge base"),
    ("Planning workflow: ", "7-step agentic state machine with conditional logic"),
    ("Memory: ", "In-memory session state (EventContext, ChatHistory, PlanningSession) across all turns"),
    ("LLM: ", "Google Gemini 2.0 Flash, with a comprehensive template-based fallback (demo mode)"),
    ("External API: ", "Spoonacular Food API for recipe-based shopping list enrichment (Tier 2)"),
    ("Frontend: ", "React + Vite + Tailwind CSS with live workflow progress tracking"),
]
for bold_part, rest in bullets_overview:
    add_bullet(rest, bold_prefix=bold_part)

add_hr()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — TESTING STRATEGY
# ══════════════════════════════════════════════════════════════════════════════
add_heading("2. Testing Strategy", 1)

# 2.1
add_heading("2.1  Manual End-to-End Testing", 2)
add_para(
    "The primary approach was scenario-based end-to-end testing — running complete planning "
    "workflows and verifying the system's behavior at each step.",
    color=TEXT_GREY, size_pt=11
)

add_table(
    ["Scenario", "Expected Behavior", "Outcome"],
    [
        ["Complete intake in one message",       "System extracts all 5 fields, skips to retrieval",                          "✅ Pass"],
        ["Intake across multiple turns",         "Context accumulates incrementally across messages",                          "✅ Pass"],
        ["Budget conflict: < $8/person",         "Conflict detection fires, halts planning, asks for resolution",             "✅ Pass"],
        ["Tight timeline: event < 7 days away",  "Timeline conflict flagged with specific actionable warning",                 "✅ Pass"],
        ["Vegetarian + gluten-free guests",      "dietary_restrictions extracted, retrieves dietary_guidelines.json",         "✅ Pass"],
        ["Children attending",                   "has_children = true, retrieves entertainment/activities content",           "✅ Pass"],
        ["Artifact generation",                  "All 3 artifacts produced (checklist, shopping list, schedule)",             "✅ Pass"],
        ["Recipe suggestions (Spoonacular on)",  "Shopping list enriched with real recipe ingredients",                       "✅ Pass"],
        ["Recipe suggestions (Spoonacular off)", "Fallback recipes with pre-defined ingredient lists used instead",           "✅ Pass"],
        ["Post-artifact adjustments",            "User can modify context and regenerate after plan is complete",             "✅ Pass"],
    ],
    col_widths=[2.4, 2.9, 0.8]
)

# 2.2
add_heading("2.2  Context Extraction Testing", 2)
add_para(
    "Context extraction was tested with various natural-language phrasings to verify that the "
    "regex patterns and Gemini-based parsing handled them correctly:",
    color=TEXT_GREY, size_pt=11
)

add_table(
    ["Input Phrase", "Field Tested", "Result"],
    [
        ['"20 guests"',            "guest_count_estimated", "✅ Correct"],
        ['"guests to 15"',         "guest_count_estimated", "✅ Correct (after regex fix)"],
        ['"$300"',                 "budget_total",          "✅ Correct"],
        ['"budget to 400"',        "budget_total",          "✅ Correct (after regex fix)"],
        ['"2026-04-25"',           "event_date (ISO)",      "✅ Correct"],
        ['"April 25"',             "event_date (written)",  "✅ Correct (after date parser added)"],
        ['"at home"',              "venue_type",            "✅ Correct"],
        ['"birthday party"',       "event_type",            "✅ Correct"],
        ['"2 vegetarians"',        "dietary_restrictions",  "✅ Correct"],
    ],
    col_widths=[1.9, 2.1, 2.1]
)

# 2.3
add_heading("2.3  RAG Retrieval Quality Testing", 2)
add_para(
    "Retrieval was tested by issuing targeted queries and verifying that semantically relevant "
    "documents ranked highest:",
    color=TEXT_GREY, size_pt=11
)

add_table(
    ["Query", "Top Retrieved Document", "Relevance Score"],
    [
        ["birthday party planning",           "birthday_party_guide.json",    "0.81"],
        ["dietary vegetarian gluten-free",    "dietary_guidelines.json",       "0.78"],
        ["shopping 20 guests food quantities","shopping_list_templates.json",  "0.75"],
        ["day of schedule timeline setup",    "day_of_schedule_samples.json",  "0.77"],
        ["children entertainment activities", "entertainment_ideas.json",      "0.72"],
        ["outdoor venue accessibility",       "accessibility_guide.json",      "0.68"],
    ],
    col_widths=[2.4, 2.5, 1.2]
)

# 2.4
add_heading("2.4  Fallback Mode (Demo Mode) Testing", 2)
add_para(
    "The system was deliberately tested without a valid GOOGLE_API_KEY to verify that:",
    color=TEXT_GREY, size_pt=11
)
for b in [
    "Template-based responses are coherent and collect missing fields correctly",
    "Context extraction using regex (not LLM) correctly parses all required fields",
    "The workflow still advances through all 7 steps to artifact generation",
    "Artifacts are generated using artifact_generator.py templates",
]:
    add_bullet(b)

add_para(
    "Fallback mode produces a functional — though less nuanced — planning experience, which is by design.",
    color=TEXT_GREY, size_pt=11, italic=True
)

# 2.5
add_heading("2.5  API Failure Testing", 2)

add_table(
    ["Failure Injected", "System Behavior"],
    [
        ["No GOOGLE_API_KEY",                     "Falls back to template responses; workflow still completes"],
        ["Invalid / expired Gemini key (404)",    "Falls back to template responses"],
        ["Gemini quota exceeded (429)",           "Falls back to template responses"],
        ["No SPOONACULAR_API_KEY",                "Uses _fallback_recipes with pre-defined ingredient lists"],
        ["Spoonacular search returns 0 results",  "Uses fallback; randomized subset prevents repetition"],
        ["Backend restart mid-session",           "Session lost (in-memory); user must start a new session"],
    ],
    col_widths=[2.8, 3.3]
)

add_hr()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — FAILURE CASES AND LIMITATIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("3. Failure Cases and Limitations", 1)

add_heading("3.1  Known Failure Cases", 2)

failures = [
    (
        "F1 — Session Loss on Backend Restart",
        "All session data (event context, chat history, workflow state, artifacts) is stored in Python "
        "process memory. If the backend server restarts or crashes, all active sessions are lost.",
        "Users must start a new planning session from scratch.",
        "PlanningSession.to_dict() / from_dict() serialization exists and SessionManager accepts a persist_dir "
        "parameter, but file-based persistence was not activated in this implementation. Enabling it would "
        "resolve this limitation."
    ),
    (
        "F2 — Context Extraction Failures in Complex Phrasing",
        "The fallback regex-based extractor handles common phrasings but can miss unusual constructions "
        "(e.g., \"we'll have roughly two dozen people\", or implied budgets like \"maybe two fifty\").",
        "In demo/fallback mode, the user may be asked to re-enter information already provided.",
        "When Gemini is available, the LLM handles these gracefully. Regex patterns were iteratively expanded "
        "to cover the most common phrasings found in testing."
    ),
    (
        "F3 — Conflict Detection Loops",
        "If a detected conflict cannot be resolved (e.g., a date that is already past), the system can get "
        "stuck asking the user to resolve it repeatedly.",
        "The planning workflow stalls at the clarification step.",
        "A bypass was added: if ctx.detected_conflicts is already set when the user sends a message, the system "
        "routes directly to planning and clears the conflict list, allowing progress even without a resolution."
    ),
    (
        "F4 — Spoonacular API Zero-Result Queries",
        "Overly specific search queries (e.g., \"birthday party vegetarian finger food for 20 guests\") "
        "frequently returned 0 results from the Spoonacular API.",
        "The enrichment step silently fails; shopping list contains no recipe suggestions.",
        "Search queries were simplified to core food terms. Random offsets were added for variety. "
        "_fallback_recipes with complete ingredient lists ensures suggestions always appear."
    ),
    (
        "F5 — Artifact Over/Undergeneration",
        "The LLM may occasionally produce minimal (too few tasks/items) or verbose artifacts depending "
        "on how event context was described.",
        "Artifacts may not reflect the full scope of the event if context is sparse.",
        "artifact_generator.py provides template-based fallback artifacts calibrated to event type, "
        "guest count, and budget to ensure a reasonable baseline."
    ),
    (
        "F6 — No Cross-Session Knowledge",
        "Each planning session is completely isolated. The system has no memory of previous users, "
        "prior events, or commonly requested event types.",
        "No personalization or learning from past sessions.",
        "By design for this implementation. Addressable with a user account system and persistent logs."
    ),
]

for title, desc, impact, mitigation in failures:
    add_heading(title, 3)
    add_para("Description: " + desc, color=TEXT_GREY, size_pt=11)
    add_para("Impact: " + impact, color=TEXT_GREY, size_pt=11)
    add_para("Mitigation: " + mitigation, color=TEXT_GREY, size_pt=11)

add_heading("3.2  System Limitations Summary", 2)

add_table(
    ["Limitation", "Notes"],
    [
        ["In-memory only",                  "Sessions lost on server restart; persist_dir option not activated"],
        ["Single concurrent user per session", "Sessions are isolated; no collaborative planning"],
        ["No real-time conflict checking",  "Conflicts checked only at conflict_detection step"],
        ["No calendar integration",         "Dates tracked as strings; no iCal / Google Calendar sync"],
        ["Gemini rate limits",              "Free tier ~1,500 req/day; heavy use triggers 429 → fallback mode"],
        ["English only",                    "Regex patterns and prompts are English-specific"],
        ["Static knowledge base",           "23 documents authored at build time; no live web updates"],
    ],
    col_widths=[2.5, 3.6]
)

add_hr()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — PROMPT AND WORKFLOW DESIGN DECISIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("4. Prompt and Workflow Design Decisions", 1)

add_heading("4.1  System Prompt Design", 2)
add_para(
    "The base system prompt (SYSTEM_PROMPT_BASE in llm_service.py) was built around four constraints:",
    color=TEXT_GREY, size_pt=11
)

prompt_decisions = [
    ("Force KB grounding",
     "The prompt instructs: \"Base ALL recommendations on the provided knowledge base context\" and "
     "\"Never fabricate costs or quantities — use the retrieved documents as your reference.\" Without "
     "this, Gemini tends to produce confident but hallucinated numbers that contradict the knowledge base."),
    ("Explicit citation requirement",
     "\"Cite your sources explicitly using the format [Source: document_name].\" Citations serve two "
     "purposes: they make the system's reasoning auditable, and they provide the frontend with structured "
     "metadata to display alongside responses."),
    ("Conflict detection as first-class behavior",
     "\"Detect and flag conflicts (e.g., budget too low for guest count, timeline too tight).\" By framing "
     "conflict detection as a core behavior, the LLM surfaces budget/timeline/dietary conflicts consistently "
     "without needing a separate prompt for each scenario type."),
    ("Context injection at every call",
     "Each LLM call includes the current EventContext as a JSON summary and the current workflow state. "
     "The model always has the full planning state, not just the current message — essential for multi-turn "
     "coherence."),
]
for bold_part, rest in prompt_decisions:
    add_bullet(rest, bold_prefix=f"{bold_part}: ")

add_heading("4.2  Workflow State Machine Design", 2)
add_para(
    "The 7-step workflow is a conditional state machine rather than a fixed linear sequence. "
    "Key design decisions:",
    color=TEXT_GREY, size_pt=11
)

workflow_decisions = [
    ("Why not a simple chain-of-thought?",
     "A single \"plan this event\" prompt has no mechanism to handle missing information, detect conflicts "
     "before generating artifacts, or update the plan when context changes. The state machine ensures each "
     "concern is handled in the correct order."),
    ("Intake → Auto-advance",
     "When all 5 required fields are collected, the system automatically advances without requiring an "
     "explicit user command. This avoids friction and keeps the conversation flowing naturally."),
    ("RAG during intake (not just retrieval step)",
     "The system retrieves KB chunks even during the intake step so early responses are grounded — "
     "e.g., if the user says \"birthday party\" but hasn't provided a budget yet, responses still "
     "reference birthday_party_guide.json."),
    ("Conflict detection gate",
     "Artifacts cannot be generated until conflicts are resolved (or overridden). This guardrail prevents "
     "producing a shopping list for a $5/person budget without flagging the issue."),
    ("Validation as a free-chat zone",
     "Between planning and artifact generation, the validation step allows open-ended questions and "
     "adjustments. It triggers a fresh RAG retrieval for each question, so answers remain grounded."),
    ("Adjustment keyword detection",
     "In validation, adjustment-intent keywords trigger a re-run of the full planning pipeline rather "
     "than just a chat reply. This handles cases like \"change the budget to $500\" — the system updates "
     "EventContext, re-retrieves, and re-generates the plan."),
]
for bold_part, rest in workflow_decisions:
    add_bullet(rest, bold_prefix=f"{bold_part}: ")

add_heading("4.3  Context Accumulation Design", 2)
add_para(
    "The EventContext object accumulates across turns rather than resetting. Every user message is "
    "processed through extract_event_context_from_intake() before routing to the workflow step handler. "
    "This means a user who provides information across three separate messages gets the same result as "
    "one who provides everything at once.",
    color=TEXT_GREY, size_pt=11
)
add_para(
    "The get_summary() method on EventContext serializes the current state into a compact string injected "
    "into every LLM call, ensuring the model always has the full planning context.",
    color=TEXT_GREY, size_pt=11
)

add_heading("4.4  Fallback Design Philosophy", 2)
add_para(
    "The fallback (demo mode) was treated as a distinct, functional mode rather than a degraded experience:",
    color=TEXT_GREY, size_pt=11
)
for b in [
    "Dynamic responses: The fallback acknowledges already-collected fields and lists only remaining missing ones — not a static boilerplate.",
    "Regex breadth: Multiple patterns cover the most common phrasings for each field, including edge cases from testing.",
    "Template artifacts: artifact_generator.py generates structured artifacts calibrated to event type, guest count, and budget.",
]:
    add_bullet(b)

add_hr()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — EVIDENCE OF RETRIEVAL GROUNDING
# ══════════════════════════════════════════════════════════════════════════════
add_heading("5. Evidence of Retrieval Grounding", 1)

add_heading("5.1  How Retrieval Grounding Works", 2)
add_para(
    "At every workflow step, the system retrieves 4–8 semantically relevant chunks from the knowledge base "
    "using the sentence-transformers model and ChromaDB vector store. These chunks are formatted into a "
    "context block injected into the LLM prompt:",
    color=TEXT_GREY, size_pt=11
)

add_code([
    "=== RETRIEVED KNOWLEDGE BASE CONTEXT ===",
    "",
    "[Source 1: Birthday Party Planning Guide | doc_id: birthday_party_guide | similarity: 81%]",
    "<chunk text>",
    "",
    "[Source 2: Event Budget Planning Guide | doc_id: budget_planning_guide | similarity: 74%]",
    "<chunk text>",
    "",
    "=== END OF RETRIEVED CONTEXT ===",
])
doc.add_paragraph()

add_heading("5.2  Query Enrichment", 2)
add_para(
    "The retrieval query is enriched with the current event context before embedding, improving relevance. "
    "For example, a query of \"shopping list\" for a vegetarian birthday party with 20 guests becomes:",
    color=TEXT_GREY, size_pt=11
)
add_code([
    'enriched_query = "shopping list birthday party home dietary vegetarian 20 guests"',
])
doc.add_paragraph()
add_para(
    "This yields higher relevance scores for shopping_list_templates.json and dietary_guidelines.json "
    "than the bare query would.",
    color=TEXT_GREY, size_pt=11, italic=True
)

add_heading("5.3  Citations Visible to the User", 2)
add_para(
    "Every response carries a citations array displayed in the UI alongside the AI message, making "
    "retrieval grounding transparent. During testing, citations consistently matched the query context:",
    color=TEXT_GREY, size_pt=11
)

add_table(
    ["Workflow Step", "Documents Retrieved"],
    [
        ["Intake",              "Birthday Party Planning Guide, Event Budget Planning Guide"],
        ["Retrieval",           "Birthday Party Planning Guide, Shopping List Templates, Dietary Guidelines, Budget Planning Guide, Day-of Schedule Samples"],
        ["Conflict Detection",  "Event Budget Planning Guide (flagged: $15/person low for sit-down meal)"],
        ["Planning",            "Birthday Party Planning Guide, Catering Guidelines, Vendor & Decoration Ideas"],
        ["Artifact Generation", "Shopping List Templates, Day-of Schedule Samples, Catering Guidelines"],
    ],
    col_widths=[1.8, 4.3]
)

add_heading("5.4  Conflict Detection Grounded in KB", 2)
add_para(
    "Conflict detection is not rule-based — the LLM receives the retrieved KB context and is asked to "
    "identify issues against knowledge-base guidelines. This means conflict messages are sourced from "
    "the knowledge base, not from hardcoded thresholds. For example:",
    color=TEXT_GREY, size_pt=11
)
for b in [
    "\"Your budget of $300 for 20 guests is $15/person. The knowledge base recommends $20–35/person for a home birthday party with a sit-down meal.\"",
    "\"The event is 12 days away. Custom bakery cakes typically require 1–2 weeks lead time.\"",
]:
    add_bullet(b)

add_para(
    "The system adapts to different event types because the retrieved chunks differ — what constitutes a "
    "budget conflict for a formal dinner party differs from a casual birthday.",
    color=TEXT_GREY, size_pt=11, italic=True
)

add_heading("5.5  Spoonacular as a Second Grounding Layer", 2)
add_para(
    "For shopping list enrichment, the Spoonacular Food API provides a second grounding layer beyond "
    "the static knowledge base:",
    color=TEXT_GREY, size_pt=11
)
for b in [
    "The system identifies the event type from EventContext",
    "A targeted recipe search query is constructed (e.g., \"birthday cake dessert\")",
    "Spoonacular returns real recipe data with ingredient quantities scaled to the guest count",
    "Ingredients are merged into the shopping list with estimated costs",
]:
    add_bullet(b)

add_heading("5.6  Sample Retrieval → Response Chain", 2)
add_para(
    "For a vegetarian birthday party with 20 guests, $300, the system retrieved the following from "
    "dietary_guidelines.json:",
    color=TEXT_GREY, size_pt=11
)
add_code([
    '"For vegetarian guests, ensure at least 30% of food options are vegetarian-friendly.',
    ' Label all items clearly at the buffet. Consider vegetarian protein sources: hummus,',
    ' cheese platters, stuffed mushrooms, vegetarian pasta."',
])
doc.add_paragraph()
add_para(
    "The resulting planning response included: \"Since you have vegetarian guests, the knowledge base "
    "recommends dedicating at least 30% of your food options to vegetarian choices — hummus and veggie "
    "platters, stuffed mushrooms, and vegetarian pasta work well for 20 guests at this budget "
    "[Source: Dietary Guidelines].\"",
    color=TEXT_GREY, size_pt=11, italic=True
)
add_para(
    "This demonstrates the full chain: RAG retrieval → context injection → grounded generation → "
    "cited response.",
    color=TEXT_GREY, size_pt=11
)

add_hr()

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX — ARCHITECTURE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
add_heading("Appendix — System Architecture Summary", 1)

add_code([
    "User Message",
    "    ↓",
    "extract_event_context_from_intake()     ← Gemini JSON extraction + regex fallback",
    "    ↓",
    "EventContext updated",
    "    ↓",
    "Route to workflow step handler",
    "    ↓",
    "RAG retrieve (sentence-transformers + ChromaDB)",
    "    enriched_query = user_message + event context fields",
    "    top-k = 4 to 8 chunks",
    "    ↓",
    "_build_rag_context() → context string for LLM prompt",
    "    ↓",
    "chat_with_context() → Gemini API call  (or template fallback)",
    "    system_prompt: SYSTEM_PROMPT_BASE + event context + RAG context",
    "    messages: last 6–10 turns of chat history",
    "    ↓",
    "Response + citations returned to frontend",
    "    ↓",
    "(At artifact_generation step):",
    "generate_artifact_json() × 3 → Task Checklist, Shopping List, Day-of Schedule",
    "    ↓",
    "(If Spoonacular enabled):",
    "enrich_shopping_list() → recipe search → ingredient scaling → merged list",
])

doc.add_paragraph()
add_para(
    "Report prepared for SE4471B Course Project, Assignment 3  |  Group 2 — Sara Daher · Aya Maree",
    color=TEXT_GREY, size_pt=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER
)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
output_path = "Reliability_and_Evaluation_Report.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
